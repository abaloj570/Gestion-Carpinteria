import tkinter as tk
from tkinter import messagebox, ttk
import os, json, pandas as pd
from datetime import datetime
from tkcalendar import DateEntry
from database_manager import *
try:
    from docx import Document
except ImportError:
    Document = None

class CRM_Carpinteria_Contable:
    def __init__(self, root):
        self.root = root
        self.root.title("Carpintería PRO v2.0 - Control de Cobros Inteligente")
        self.root.geometry("1600x900")
        self.root.configure(bg="#f0f2f5")
        
        inicializar_db()
        realizar_backup()
        self.filtro_estado = "TODOS"

        # --- CABECERA ---
        header = tk.Frame(root, bg="#2c3e50", pady=10)
        header.pack(fill="x")
        tk.Label(header, text="🔍 BUSCAR:", bg="#2c3e50", fg="white", font=("Arial", 10, "bold")).pack(side="left", padx=15)
        self.ent_buscar = tk.Entry(header, width=25, font=("Arial", 11))
        self.ent_buscar.pack(side="left", padx=5)
        self.ent_buscar.bind("<KeyRelease>", lambda e: self.actualizar_tabla())

        # (Resto de botones de filtro iguales...)
        filtros_f = tk.Frame(header, bg="#2c3e50")
        filtros_f.pack(side="right", padx=20)
        for t, c in [("TODOS","#ecf0f1"), ("TALLER","#3498db"), ("MONTAJE","#9b59b6"), ("FINALIZADA","#27ae60")]:
            tk.Button(filtros_f, text=t, bg=c, font=("Arial", 8, "bold"), command=lambda x=t: self.set_filtro(x)).pack(side="left", padx=2)

        # --- ALTA ---
        input_f = tk.LabelFrame(root, text=" Registro de Obra ", bg="#f0f2f5", pady=10, padx=15)
        input_f.pack(fill="x", padx=15, pady=5)
        tk.Label(input_f, text="Obra:").grid(row=0, column=0)
        self.ent_obra = tk.Entry(input_f, width=20); self.ent_obra.grid(row=0, column=1, padx=5)
        tk.Label(input_f, text="Cliente:").grid(row=0, column=2)
        self.ent_empresa = tk.Entry(input_f, width=20); self.ent_empresa.grid(row=0, column=3, padx=5)
        tk.Label(input_f, text="Entrega:").grid(row=0, column=4)
        self.cal_ent = DateEntry(input_f, width=12, date_pattern='dd/mm/yyyy'); self.cal_ent.grid(row=0, column=5, padx=5)
        tk.Button(input_f, text="➕ REGISTRAR", command=self.agregar, bg="#2ecc71", font=("Arial", 9, "bold")).grid(row=0, column=6, padx=20)

        # --- TABLA ---
        cols = ("⭐", "Obra", "Cliente", "Entrega", "Estado", "📏 Corte", "📂 Presp.", "💶 Fact.", "📸 Fotos", "💸 Pago")
        self.tabla = ttk.Treeview(root, columns=cols, show='headings', height=25)
        for c in cols: 
            self.tabla.heading(c, text=c.upper())
            self.tabla.column(c, width=90, anchor="center")
        self.tabla.column("Obra", width=200, anchor="w")
        self.tabla.column("Cliente", width=180, anchor="w")
        self.tabla.column("⭐", width=35)
        
        self.tabla.tag_configure('prioridad', background='#ff9f43', foreground='white')
        self.tabla.tag_configure('atrasado', background='#ee5253', foreground='white')
        self.tabla.tag_configure('urgente', background='#feca57')
        self.tabla.pack(expand=True, fill="both", padx=15, pady=10)
        self.tabla.bind("<Double-1>", self.gestionar_click)

        # --- PANEL INFERIOR ---
        bottom_f = tk.Frame(root, bg="#f0f2f5")
        bottom_f.pack(fill="x", padx=15, pady=10)
        tk.Button(bottom_f, text="📄 EXPORTAR A WORD", command=self.exportar_docx, bg="#2980b9", fg="white", font=("Arial", 9, "bold")).pack(side="left")
        tk.Button(bottom_f, text="🗑️ ELIMINAR", command=self.borrar, bg="#d63031", fg="white", font=("Arial", 9, "bold")).pack(side="right")

        self.actualizar_tabla()

    def set_filtro(self, st):
        self.filtro_estado = st
        self.actualizar_tabla()

    def agregar(self):
        o, e, f = self.ent_obra.get().strip(), self.ent_empresa.get().strip(), self.cal_ent.get()
        if o:
            ruta = crear_carpetas_obra(o)
            guardar_registro({'Obra': o, 'Empresa': e, 'Ruta_Carpeta': ruta, 'Estado': 'PRESUPUESTO', 'Fecha_Entrega': f})
            self.actualizar_tabla(); self.ent_obra.delete(0, tk.END); self.ent_empresa.delete(0, tk.END)

    def abrir_seguro(self, ruta_base, sub=None):
        target = os.path.join(ruta_base, sub) if sub else ruta_base
        if not os.path.exists(target): os.makedirs(target, exist_ok=True)
        os.startfile(target)

    def actualizar_tabla(self):
        for i in self.tabla.get_children(): self.tabla.delete(i)
        df, hoy, t = obtener_todo(), datetime.now().date(), self.ent_buscar.get().lower()
        for idx, row in df.iterrows():
            if (t in str(row['Obra']).lower() or t in str(row['Empresa']).lower()) and \
               (self.filtro_estado == "TODOS" or row['Estado'] == self.filtro_estado):
                tag, f_str = '', str(row.get('Fecha_Entrega', ''))
                prio = "⭐" if row.get('Prioridad') == 'SI' else ""
                pago = "✅" if row.get('Cobro') == 'PAGADO' else "❌"
                try:
                    f_ent = datetime.strptime(f_str, '%d/%m/%Y').date()
                    dif = (f_ent - hoy).days
                    if row['Estado'] != "FINALIZADA":
                        if row.get('Prioridad') == 'SI': tag = 'prioridad'
                        elif dif < 0: tag = 'atrasado'
                        elif dif <= 3: tag = 'urgente'
                except: pass
                self.tabla.insert("", "end", values=[prio, row['Obra'], row['Empresa'], f_str, row['Estado'], "📏", "📂", "💶", "📸", pago], iid=idx, tags=(tag,))

    def gestionar_click(self, event):
        item_id = self.tabla.focus()
        if not item_id: return
        col, idx = self.tabla.identify_column(event.x), int(item_id)
        row = obtener_todo().iloc[idx]
        r = row['Ruta_Carpeta']
        
        if col in ("#1", "#2"): self.ventana_detalles(idx, row)
        elif col == "#3": self.abrir_seguro(r)
        elif col == "#6": self.abrir_seguro(r, "Despiece")
        elif col == "#7": self.abrir_seguro(r, "Presupuestos")
        elif col == "#8": self.abrir_seguro(r, "Facturas")
        elif col == "#9": self.abrir_seguro(r, "Fotos")
        elif col == "#10":
            nuevo = 'PAGADO' if row.get('Cobro') != 'PAGADO' else 'PENDIENTE'
            actualizar_fila(idx, {'Cobro': nuevo})
            self.actualizar_tabla()

    def ventana_detalles(self, idx_real, datos):
        win = tk.Toplevel(self.root); win.title(f"Detalles: {datos['Obra']}"); win.geometry("950x850")
        nb = ttk.Notebook(win); nb.pack(expand=True, fill="both", padx=10, pady=10)

        # --- PESTAÑA 1: ESTADO Y COBROS ---
        f1 = tk.Frame(nb); nb.add(f1, text="🏗️ ESTADO Y FINANZAS")
        
        # Sección Prioridad y Estado
        top_frame = tk.LabelFrame(f1, text=" Producción ", padx=10, pady=10)
        top_frame.pack(fill="x", padx=10, pady=5)
        
        def toggle_prio():
            nueva = 'SI' if datos.get('Prioridad') != 'SI' else 'NO'
            actualizar_fila(idx_real, {'Prioridad': nueva}); self.actualizar_tabla(); win.destroy()
        
        tk.Button(top_frame, text="⭐ ALTERNAR PRIORIDAD", command=toggle_prio, bg="#f39c12", fg="white").grid(row=0, column=0, padx=5)
        
        tk.Label(top_frame, text="Fase:").grid(row=0, column=1, padx=5)
        combo = ttk.Combobox(top_frame, values=["PRESUPUESTO", "TALLER", "MONTAJE", "FINALIZADA"], state="readonly")
        combo.set(datos['Estado']); combo.grid(row=0, column=2, padx=5)

        # --- NUEVA SECCIÓN: AUDITORÍA DE FACTURAS ---
        fact_frame = tk.LabelFrame(f1, text=" 💰 COMPROBACIÓN DE FACTURAS ", padx=10, pady=10, fg="#2c3e50")
        fact_frame.pack(fill="both", expand=True, padx=10, pady=5)

        lista_impagos = tk.Listbox(fact_frame, height=5, font=("Arial", 10), fg="#d63031")
        lista_impagos.pack(fill="x", pady=5)

        lbl_status = tk.Label(fact_frame, text="Estado: No verificado", font=("Arial", 9, "italic"))
        lbl_status.pack()

        def verificar_impagados():
            ruta_facturas = os.path.join(datos['Ruta_Carpeta'], "Facturas")
            lista_impagos.delete(0, tk.END)
            pendientes = []
            
            if os.path.exists(ruta_facturas):
                archivos = os.listdir(ruta_facturas)
                for f in archivos:
                    # Regla: Si el nombre tiene PENDIENTE o IMPAGADO
                    if "PENDIENTE" in f.upper() or "IMPAGADO" in f.upper():
                        pendientes.append(f)
                
                if pendientes:
                    for p in pendientes:
                        lista_impagos.insert(tk.END, f"⚠️ {p}")
                    lbl_status.config(text=f"Se han encontrado {len(pendientes)} facturas sin pagar.", fg="#d63031")
                else:
                    lbl_status.config(text="✅ Todo al día. No hay archivos marcados como pendientes.", fg="#27ae60")
            else:
                lbl_status.config(text="❌ Carpeta de facturas no encontrada.")

        tk.Button(fact_frame, text="🔍 ESCANEAR CARPETA DE FACTURAS", command=verificar_impagados, bg="#34495e", fg="white").pack(pady=5)

        def guardar_todo():
            actualizar_fila(idx_real, {'Estado': combo.get()})
            self.actualizar_tabla(); win.destroy()
            
        tk.Button(f1, text="💾 GUARDAR CAMBIOS", command=guardar_todo, bg="#2980b9", fg="white", font=("Arial", 10, "bold"), height=2).pack(pady=10)

        # --- PESTAÑA 2: DIARIO DE NOTAS ---
        f2 = tk.Frame(nb); nb.add(f2, text="📝 DIARIO DE NOTAS")
        t_h = ttk.Treeview(f2, columns=("Fecha","Nota"), show='headings', height=18)
        t_h.heading("Fecha", text="FECHA/HORA"); t_h.column("Fecha", width=120)
        t_h.heading("Nota", text="ANOTACIÓN"); t_h.column("Nota", width=600)
        t_h.pack(fill="both", expand=True, padx=10, pady=10)
        
        lh = json.loads(datos['Historial_Fechas']) if pd.notna(datos['Historial_Fechas']) else []
        for h in lh: t_h.insert("", "end", values=(h['f'], h['n']))

        entry_f = tk.Frame(f2); entry_f.pack(fill="x", padx=10, pady=10)
        self.e_nota = tk.Entry(entry_f, font=("Arial", 11))
        self.e_nota.pack(side="left", fill="x", expand=True, padx=5)
        
        def add_nota():
            txt = self.e_nota.get().strip()
            if txt:
                nueva = {'f': datetime.now().strftime("%d/%m %H:%M"), 'n': txt}
                lh.append(nueva)
                actualizar_fila(idx_real, {'Historial_Fechas': json.dumps(lh)})
                t_h.insert("", "end", values=(nueva['f'], nueva['n']))
                self.e_nota.delete(0, tk.END); self.actualizar_tabla()

        tk.Button(entry_f, text="➕ AÑADIR NOTA", command=add_nota, bg="#2ecc71", fg="white").pack(side="right")

    def exportar_docx(self):
        if not Document:
            messagebox.showerror("Error", "Instala: pip install python-docx")
            return
        sel = self.tabla.selection()
        if not sel: return
        idx = int(sel[0])
        row = obtener_todo().iloc[idx]
        doc = Document()
        doc.add_heading(f"OBRA: {row['Obra']}", 0)
        doc.add_paragraph(f"Cliente: {row['Empresa']}")
        doc.add_heading("NOTAS", level=1)
        lh = json.loads(row['Historial_Fechas']) if pd.notna(row['Historial_Fechas']) else []
        for h in lh: doc.add_paragraph(f"[{h['f']}] {h['n']}", style='List Bullet')
        doc.save(f"Ficha_{row['Obra']}.docx")
        os.startfile(f"Ficha_{row['Obra']}.docx")

    def borrar(self):
        sel = self.tabla.selection()
        if not sel: return
        n = self.tabla.item(sel[0])['values'][1]
        if messagebox.askyesno("Confirmar", f"¿Eliminar obra '{n}'?"):
            if borrar_obra(n): self.actualizar_tabla()

if __name__ == "__main__":
    root = tk.Tk(); app = CRM_Carpinteria_Contable(root); root.mainloop()